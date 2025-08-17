# pyhton-docx로 DOCX 생성
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.application.ports import DocumentRendererPort
from app.application.dto import GenerateCommand, GenerateResult
import io

class DocxRenderer(DocumentRendererPort):
    def render(self, cmd: GenerateCommand) -> GenerateResult:
        ast = cmd.ast
        doc = Document()

        if ast.title:
            p = doc.add_paragraph()
            run = p.add_run(ast.title)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for s in ast.sections:
            if s.heading:
                ph = doc.add_paragraph()
                rh = ph.add_run(s.heading)
                rh.bold = True
                rh.font.size = Pt(12)
            if s.body:
                pb = doc.add_paragraph(s.body)
                for t in (s.tables or []):
                    table = doc.add_table(rows=1, cols=len(t.headers))
                    hdr = table.rows[0].cells
                    for j, h in enumerate(t.headers):
                        hdr[j].text = h
                    for row in t.rows:
                        cells = table.add_row().cells
                        for j, val in enumerate(row):
                            cells[j].text = str(val)

        if ast.footer:
            section = doc.sections[-1]
            footer = section.footer
            par = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            run = par.add_run(ast.footer)
            run.font.size = Pt(9)

        bio = io.BytesIO()
        doc.save(bio)
        return GenerateResult(
            filename="document.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            data=bio.getvalue(),
        )